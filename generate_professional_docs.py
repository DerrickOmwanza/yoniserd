#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generate Professional Business Documents for YoNISeRD Website Project
Creates: Invoice 2, Updated Project Quotation, Updated Website Development Agreement
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime, timedelta

def add_header(doc, title, subtitle=None):
    """Add professional header to document"""
    # Add title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.runs[0]
    title_run.font.color.rgb = RGBColor(5, 15, 42)  # Primary dark color
    
    # Add subtitle if provided
    if subtitle:
        subtitle_para = doc.add_paragraph(subtitle)
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle_para.runs[0]
        subtitle_run.font.italic = True
        subtitle_run.font.color.rgb = RGBColor(126, 187, 191)  # Primary blue
    
    doc.add_paragraph()  # Spacing

def add_section_title(doc, title):
    """Add section title"""
    para = doc.add_heading(title, level=1)
    para_format = para.paragraph_format
    para_format.space_before = Pt(12)
    para_format.space_after = Pt(6)
    para.runs[0].font.color.rgb = RGBColor(5, 15, 42)

def create_invoice_v2():
    """Create Invoice Document 2"""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # Header
    add_header(doc, "INVOICE", "Professional Website Development Services")
    
    # Invoice details
    details = [
        ("Invoice Number:", "INV-002-2026"),
        ("Invoice Date:", "01st February 2026"),
        ("Due Date:", "01st March 2026"),
        ("Invoice Status:", "Ready for Payment"),
    ]
    
    details_table = doc.add_table(rows=len(details), cols=2)
    details_table.autofit = False
    details_table.allow_autofit = False
    
    for i, (label, value) in enumerate(details):
        details_table.rows[i].cells[0].text = label
        details_table.rows[i].cells[1].text = value
        details_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    # Bill To / From
    bill_section = doc.add_heading("BILLING INFORMATION", level=1)
    bill_section.runs[0].font.color.rgb = RGBColor(5, 15, 42)
    
    bill_table = doc.add_table(rows=1, cols=2)
    bill_table.autofit = False
    
    # Bill To
    bill_to_cell = bill_table.rows[0].cells[0]
    bill_to_cell.text = ""
    p1 = bill_to_cell.paragraphs[0]
    p1.add_run("BILL TO:\n").font.bold = True
    bill_to_cell.add_paragraph("Youth Network Integrated Services for Research and Development")
    bill_to_cell.add_paragraph("(YoNISeRD)")
    bill_to_cell.add_paragraph("Kisii, Kenya")
    bill_to_cell.add_paragraph("Email: [Client Email]")
    bill_to_cell.add_paragraph("Phone: [Client Phone]")
    
    # From
    from_cell = bill_table.rows[0].cells[1]
    from_cell.text = ""
    p2 = from_cell.paragraphs[0]
    p2.add_run("FROM:\n").font.bold = True
    from_cell.add_paragraph("Derrick Omwanza")
    from_cell.add_paragraph("Professional Web Developer")
    from_cell.add_paragraph("Nairobi, Kenya")
    from_cell.add_paragraph("Email: derrick@example.com")
    from_cell.add_paragraph("Phone: +254 XXX XXX XXX")
    
    doc.add_paragraph()
    
    # Services Rendered
    add_section_title(doc, "SERVICES RENDERED")
    
    services_table = doc.add_table(rows=8, cols=4)
    services_table.style = 'Light Grid Accent 1'
    
    # Header
    header_cells = services_table.rows[0].cells
    header_cells[0].text = "Description"
    header_cells[1].text = "Quantity"
    header_cells[2].text = "Unit Rate (KSh)"
    header_cells[3].text = "Amount (KSh)"
    
    for cell in header_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Service rows
    services = [
        ("Website Design & Development\n(8 Pages: Home, About, Programs, Our Work, Gallery, News, Impact, Contact)\nCustom React 19 Application with TailwindCSS\nFully Responsive (Mobile, Tablet, Desktop)\nOptimized Performance & Accessibility (WCAG 2.1 AA)", "1", "10,000", "10,000"),
        ("Component Development & Features\n(Navigation System, Slideshows, Forms, Metrics Animations, Story Modal)\nErrorBoundary Implementation\nSEO Optimization (Meta Tags, Open Graph, Schema.org)", "1", "3,000", "3,000"),
        ("Quality Assurance & Testing\n(Comprehensive QA Audit - 100/100 Score)\nAll Pages, Buttons, Forms Verified\nResponsive Design Testing (Mobile/Tablet/Desktop)\nSlideshow Performance Optimization", "1", "2,000", "2,000"),
        ("Configuration & Deployment\n(Netlify Hosting Setup, Domain Configuration)\nForm Integration, Email Notifications\nSSL Certificate Configuration\nSecurity Setup & ErrorBoundary", "1", "1,500", "1,500"),
    ]
    
    for i, (desc, qty, rate, amount) in enumerate(services, 1):
        row = services_table.rows[i]
        row.cells[0].text = desc
        row.cells[1].text = qty
        row.cells[2].text = rate
        row.cells[3].text = amount
        
        for cell in row.cells[1:]:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Totals
    totals_row = services_table.rows[5]
    totals_row.cells[0].text = "SUBTOTAL"
    totals_row.cells[0].paragraphs[0].runs[0].font.bold = True
    totals_row.cells[3].text = "16,500"
    totals_row.cells[3].paragraphs[0].runs[0].font.bold = True
    
    tax_row = services_table.rows[6]
    tax_row.cells[0].text = "VAT (0% - As per Contract)"
    tax_row.cells[3].text = "0"
    tax_row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    total_row = services_table.rows[7]
    total_row.cells[0].text = "TOTAL DUE"
    total_row.cells[0].paragraphs[0].runs[0].font.bold = True
    total_row.cells[0].paragraphs[0].runs[0].font.size = Pt(12)
    total_row.cells[3].text = "KSh 16,500"
    total_row.cells[3].paragraphs[0].runs[0].font.bold = True
    total_row.cells[3].paragraphs[0].runs[0].font.size = Pt(12)
    
    doc.add_paragraph()
    
    # Payment Terms
    add_section_title(doc, "PAYMENT TERMS & SCHEDULE")
    
    payment_table = doc.add_table(rows=3, cols=3)
    payment_table.style = 'Light Grid Accent 1'
    
    header = payment_table.rows[0].cells
    header[0].text = "Milestone"
    header[1].text = "Amount (KSh)"
    header[2].text = "Status"
    for cell in header:
        cell.paragraphs[0].runs[0].font.bold = True
    
    payment_table.rows[1].cells[0].text = "Initial Deposit (50%)"
    payment_table.rows[1].cells[1].text = "8,250"
    payment_table.rows[1].cells[2].text = "[Paid on: 28th Jan 2026]"
    
    payment_table.rows[2].cells[0].text = "Final Payment (50%)"
    payment_table.rows[2].cells[1].text = "8,250"
    payment_table.rows[2].cells[2].text = "Due upon completion"
    
    doc.add_paragraph()
    
    # Payment Methods
    add_section_title(doc, "PAYMENT INSTRUCTIONS")
    
    doc.add_paragraph("Please remit payment using one of the following methods:")
    doc.add_paragraph()
    
    methods = [
        ("M-PESA", [
            "Business Name: Derrick Omwanza",
            "Phone: +254 XXX XXX XXX",
            "Till Number: [Till Number]"
        ]),
        ("Bank Transfer", [
            "Bank: [Bank Name]",
            "Account Name: Derrick Omwanza",
            "Account Number: [Account Number]",
            "Branch: [Branch]",
            "Swift Code: [Swift Code]"
        ])
    ]
    
    for method_name, details in methods:
        para = doc.add_paragraph()
        para.add_run(f"{method_name}: ").font.bold = True
        for detail in details:
            doc.add_paragraph(detail, style='List Bullet')
    
    doc.add_paragraph()
    
    # Deliverables Confirmed
    add_section_title(doc, "DELIVERABLES CONFIRMED")
    
    deliverables = [
        "Home page - Hero banner, impact statistics, program highlights, latest stories, CTA sections",
        "About page - Mission/Vision, CEO message, team profiles, organizational timeline, awards section",
        "Programs page - 5 program categories with detailed descriptions and partnership opportunities",
        "Our Work page - Gallery slideshow with 9 featured community initiatives and impact stories",
        "Gallery page - Standalone image slideshow with responsive design and touch support",
        "News & Stories page - Dynamic filtering by category with story modal for detailed views",
        "Impact page - Animated metrics counters, testimonials, timeline, and photo showcase",
        "Contact page - Functional form with client-side validation, volunteer CTA, multiple contact options",
        "Responsive Navigation - Desktop dropdowns (3 menus × 3-4 items), mobile hamburger menu",
        "Footer - Social media links (LinkedIn, Twitter, Facebook, Instagram) and quick links",
        "Image Slideshow Component - Auto-rotation, manual controls, touch/swipe support, responsive",
        "Impact Metrics Component - Scroll-triggered animations, responsive counters",
        "Story Modal Component - Detailed story view with smooth transitions",
        "ErrorBoundary Component - Graceful error handling and recovery",
        "Full Accessibility Implementation - WCAG 2.1 Level AA compliant (skip links, ARIA labels, keyboard navigation, focus states)",
        "SEO Optimization - Meta tags, Open Graph protocol, Twitter Cards, Schema.org structured data",
        "Performance Optimization - CSS variable consolidation, GPU acceleration, will-change properties, CSS containment",
        "Comprehensive QA Audit - 100/100 quality score, all pages and buttons verified, responsive testing, performance metrics",
        "Netlify Deployment Configuration - SPA routing, form integration, continuous deployment setup",
        "Source Code Handover - Complete React project with documentation and setup instructions",
    ]
    
    for deliverable in deliverables:
        doc.add_paragraph(deliverable, style='List Bullet')
    
    doc.add_paragraph()
    
    # Project Status
    add_section_title(doc, "PROJECT STATUS")
    
    status_para = doc.add_paragraph()
    status_para.add_run("Status: ").font.bold = True
    status_para.add_run("COMPLETE & READY FOR DEPLOYMENT").font.color.rgb = RGBColor(0, 128, 0)
    status_para.runs[-1].font.bold = True
    
    doc.add_paragraph()
    doc.add_paragraph("The website has been fully developed, thoroughly tested (comprehensive QA audit completed with 100/100 score), and optimized for production. All pages, buttons, forms, and features are verified and ready for deployment.")
    
    doc.add_paragraph()
    
    # Terms & Conditions
    add_section_title(doc, "TERMS & CONDITIONS")
    
    terms = [
        "Payment of 50% deposit is required to commence work; final 50% is due upon project completion and handover.",
        "This invoice is valid for 30 days from the date of issue.",
        "Late payment charges: 5% per month on outstanding amounts.",
        "Client is responsible for providing all final text, images, branding assets, and content.",
        "Post-launch support: 30 days of free bug fixing and minor adjustments included.",
        "Any additional features or major changes beyond the agreed scope will be billed separately at KSh 1,500 per day.",
        "The Developer retains the right to display the completed project in professional portfolio.",
    ]
    
    for i, term in enumerate(terms, 1):
        doc.add_paragraph(term, style='List Number')
    
    doc.add_paragraph()
    
    # Footer
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.add_run("Thank you for your business!")
    footer_para.runs[0].font.italic = True
    
    doc.add_paragraph()
    footer_para2 = doc.add_paragraph("Derrick Omwanza | Professional Web Developer")
    footer_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para2.runs[0].font.size = Pt(9)
    footer_para2.runs[0].font.color.rgb = RGBColor(126, 187, 191)
    
    return doc

def create_project_quotation_v2():
    """Create Updated Project Quotation"""
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # Header
    add_header(doc, "PROJECT QUOTATION", "YoNISeRD Website Development")
    
    # Quote Details
    quote_details = [
        ("Quote Reference:", "YNIS-WEB-002-2026"),
        ("Date Issued:", "01st February 2026"),
        ("Valid Until:", "01st March 2026"),
        ("Quotation Status:", "Final & Complete"),
    ]
    
    details_table = doc.add_table(rows=len(quote_details), cols=2)
    for i, (label, value) in enumerate(quote_details):
        details_table.rows[i].cells[0].text = label
        details_table.rows[i].cells[1].text = value
        details_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    # Client & Provider
    client_section = doc.add_heading("CLIENT & PROVIDER INFORMATION", level=1)
    client_section.runs[0].font.color.rgb = RGBColor(5, 15, 42)
    
    client_table = doc.add_table(rows=1, cols=2)
    
    client_cell = client_table.rows[0].cells[0]
    client_cell.text = ""
    p1 = client_cell.paragraphs[0]
    p1.add_run("CLIENT:\n").font.bold = True
    client_cell.add_paragraph("Youth Network Integrated Services for Research and Development (YoNISeRD)")
    client_cell.add_paragraph("Organization Focus: Youth Research & Development")
    client_cell.add_paragraph("Location: Kisii, Kenya")
    client_cell.add_paragraph("Contact: [Contact Name]")
    client_cell.add_paragraph("Email: [Client Email]")
    
    provider_cell = client_table.rows[0].cells[1]
    provider_cell.text = ""
    p2 = provider_cell.paragraphs[0]
    p2.add_run("PROVIDER:\n").font.bold = True
    provider_cell.add_paragraph("Derrick Omwanza")
    provider_cell.add_paragraph("Professional Web Developer & Software Engineer")
    provider_cell.add_paragraph("Specialization: React.js, Web Development, Digital Solutions")
    provider_cell.add_paragraph("Email: derrick@example.com")
    provider_cell.add_paragraph("Phone: +254 XXX XXX XXX")
    
    doc.add_paragraph()
    
    # Executive Summary
    add_section_title(doc, "PROJECT OVERVIEW")
    
    overview_text = """This quotation outlines the complete design, development, and deployment of a professional, fully-responsive website for YoNISeRD. The website will showcase the organization's mission, programs, impact, and provide multiple avenues for community engagement and partnership.

The project is built on modern web technologies (React 19, TailwindCSS, React Router 7) ensuring optimal performance, security, accessibility, and scalability. The website includes 8 fully-functional pages, responsive navigation, integrated forms, optimized slideshows, and comprehensive accessibility features (WCAG 2.1 Level AA).

All development, testing, optimization, and deployment are included in this quotation."""
    
    doc.add_paragraph(overview_text)
    
    doc.add_paragraph()
    
    # Detailed Scope
    add_section_title(doc, "DETAILED SCOPE OF WORK")
    
    doc.add_heading("1. Website Pages & Functionality", level=2)
    
    pages = {
        "Home Page": [
            "Hero banner section with compelling call-to-action",
            "Impact statistics with animated counters (7+ years, 50 communities, 75% employment)",
            "Program highlights grid (3 featured programs with descriptions)",
            "Latest stories section with link to full stories page",
            "Core values display (4 key organizational principles)",
            "Final CTA section with dual action buttons"
        ],
        "About Page": [
            "Mission & Vision statements (side-by-side cards)",
            "Organization details (founding year, legal status, registration number)",
            "CEO message with professional portrait",
            "Leadership team profiles (5 team members with photos)",
            "Organizational journey timeline (6 key milestones)",
            "Awards & recognition section (3+ achievements)",
            "Testimonials from beneficiaries and partners",
            "Core values grid (4-column responsive layout)"
        ],
        "Programs Page": [
            "Comprehensive program overview",
            "5 flagship programs with detailed descriptions",
            "Each program includes: title, description, target audience, expected outcomes",
            "Partnership opportunities section",
            "Call-to-action for program participation"
        ],
        "Our Work Page": [
            "Gallery of 9 featured community initiatives",
            "Story cards with tags (Our Work, Impact, News)",
            "Link to detailed story views",
            "Filter capability by program/category",
            "Image showcase with slideshow functionality"
        ],
        "Gallery Page": [
            "Standalone image gallery with 6+ photos",
            "Auto-rotating slideshow functionality",
            "Manual navigation (previous/next buttons)",
            "Image descriptions and captions",
            "Touch/swipe gesture support on mobile"
        ],
        "News & Stories Page": [
            "8+ story items with category filtering",
            "Story modal for detailed view of each story",
            "Category tags (Our Work, Impact, News)",
            "Filtering by category or search",
            "Publication dates and author information"
        ],
        "Impact Page": [
            "Animated impact metrics (scroll-triggered counters)",
            "Key statistics visualization (youth reached, communities, programs)",
            "Testimonials from program beneficiaries",
            "Photo timeline or showcase",
            "Data-driven impact narrative"
        ],
        "Contact Page": [
            "Contact form with 4 fields (Name, Email, Subject, Message)",
            "Client-side validation with error messages",
            "Netlify Forms integration with email notifications",
            "Contact information display (email, phone, address, social links)",
            "Office location (Google Maps integration optional)",
            "Volunteer CTA and partnership inquiry section"
        ]
    }
    
    for page_name, features in pages.items():
        para = doc.add_paragraph()
        para.add_run(f"{page_name}: ").font.bold = True
        for feature in features:
            doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_paragraph()
    
    doc.add_heading("2. Technical Features & Components", level=2)
    
    features = {
        "Navigation System": [
            "Sticky header with logo and organization name",
            "Desktop navigation: 3 dropdown menus (About Us, What We Do, Get Involved) with 3-4 items each",
            "Mobile navigation: Hamburger menu with slide-in drawer",
            "All navigation items fully functional with smooth transitions",
            "Responsive design (hidden on desktop <1024px, visible on mobile)"
        ],
        "Forms & Validation": [
            "Contact form with client-side validation",
            "Name field (required, text validation)",
            "Email field (required, regex validation for valid email format)",
            "Subject field (required, text input)",
            "Message field (required, minimum 10 characters)",
            "Error messages displayed inline with clear language",
            "Success message with animation after submission",
            "Form resets after successful submission"
        ],
        "Slideshows & Carousels": [
            "Image slideshow component with auto-rotation",
            "Manual navigation controls (previous/next buttons)",
            "Touch/swipe gesture support for mobile users",
            "Visual indicators (dots) showing current slide",
            "Image descriptions and captions",
            "Responsive aspect ratio (16:9) preventing layout shifts",
            "GPU-accelerated animations (smooth 60fps)"
        ],
        "Animations & Interactions": [
            "Impact metrics with scroll-triggered counter animations",
            "Smooth hover effects on cards and buttons",
            "Fade-in effects for images and sections",
            "Slide-up animations for content reveal",
            "Transition animations for navigation menu",
            "Respects user's reduced motion preferences (accessibility)"
        ],
        "Accessibility Features": [
            "Skip link for keyboard users (press Tab to reveal)",
            "WCAG 2.1 Level AA compliance",
            "Proper heading hierarchy (h1 > h2 > h3, no skipped levels)",
            "Color contrast ≥ 4.5:1 on all text elements",
            "Keyboard navigation enabled (Tab, Enter, Escape keys)",
            "Focus visible styling (blue outline on all interactive elements)",
            "ARIA labels on buttons and navigation items",
            "Form labels properly associated with inputs",
            "Semantic HTML structure throughout"
        ],
        "Search Engine Optimization": [
            "Meta tags for all pages (title, description, keywords)",
            "Open Graph protocol for social media sharing",
            "Twitter Card implementation",
            "Schema.org structured data for organizational information",
            "Proper heading structure for SEO",
            "Image alt text on all images",
            "Mobile-friendly design (mobile-first responsive)"
        ]
    }
    
    for feature_name, feature_list in features.items():
        para = doc.add_paragraph()
        para.add_run(f"{feature_name}: ").font.bold = True
        for feature in feature_list:
            doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_paragraph()
    
    # Investment Breakdown
    add_section_title(doc, "INVESTMENT BREAKDOWN")
    
    investment_table = doc.add_table(rows=6, cols=3)
    investment_table.style = 'Light Grid Accent 1'
    
    header = investment_table.rows[0].cells
    header[0].text = "Service Description"
    header[1].text = "Details"
    header[2].text = "Cost (KSh)"
    for cell in header:
        cell.paragraphs[0].runs[0].font.bold = True
    
    rows_data = [
        ("Website Design & Development", "8 Pages, Responsive Design, All Features", "10,000"),
        ("Component Development & Features", "Navigation, Forms, Slideshows, Animations, Components", "3,000"),
        ("Quality Assurance & Testing", "Comprehensive QA Audit, 100/100 Score, Performance Testing", "2,000"),
        ("Configuration & Deployment", "Hosting Setup, Form Integration, Security", "1,500"),
    ]
    
    for i, (service, details, cost) in enumerate(rows_data, 1):
        row = investment_table.rows[i]
        row.cells[0].text = service
        row.cells[1].text = details
        row.cells[2].text = cost
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Total row
    total_row = investment_table.rows[5]
    total_cell = total_row.cells[0]
    total_cell.text = "TOTAL PROJECT INVESTMENT"
    total_cell.paragraphs[0].runs[0].font.bold = True
    total_cell.paragraphs[0].runs[0].font.size = Pt(12)
    total_row.cells[2].text = "KSh 16,500"
    total_row.cells[2].paragraphs[0].runs[0].font.bold = True
    total_row.cells[2].paragraphs[0].runs[0].font.size = Pt(12)
    total_row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_paragraph()
    
    # Payment Schedule
    add_section_title(doc, "PAYMENT SCHEDULE")
    
    schedule_table = doc.add_table(rows=3, cols=3)
    schedule_table.style = 'Light Grid Accent 1'
    
    header = schedule_table.rows[0].cells
    header[0].text = "Phase"
    header[1].text = "Amount (KSh)"
    header[2].text = "Condition"
    for cell in header:
        cell.paragraphs[0].runs[0].font.bold = True
    
    schedule_table.rows[1].cells[0].text = "Initial Deposit"
    schedule_table.rows[1].cells[1].text = "8,250 (50%)"
    schedule_table.rows[1].cells[2].text = "Upon signing agreement"
    
    schedule_table.rows[2].cells[0].text = "Final Payment"
    schedule_table.rows[2].cells[1].text = "8,250 (50%)"
    schedule_table.rows[2].cells[2].text = "Upon completion & deployment"
    
    doc.add_paragraph()
    
    # Project Timeline
    add_section_title(doc, "PROJECT TIMELINE")
    
    timeline = [
        ("Design & Prototype", "Completed"),
        ("Development Phase", "Completed"),
        ("Quality Assurance & Testing", "Completed"),
        ("Optimization & Refinement", "Completed"),
        ("Launch & Deployment", "Ready"),
    ]
    
    timeline_table = doc.add_table(rows=len(timeline) + 1, cols=2)
    timeline_table.style = 'Light Grid Accent 1'
    
    header = timeline_table.rows[0].cells
    header[0].text = "Phase"
    header[1].text = "Status"
    for cell in header:
        cell.paragraphs[0].runs[0].font.bold = True
    
    for i, (phase, status) in enumerate(timeline, 1):
        timeline_table.rows[i].cells[0].text = phase
        timeline_table.rows[i].cells[1].text = status
    
    doc.add_paragraph()
    
    # Terms & Conditions
    add_section_title(doc, "TERMS & CONDITIONS")
    
    terms = [
        "This quotation is valid for 30 days from the date of issue.",
        "50% deposit required to commence finalization work; 50% due upon completion and successful deployment.",
        "Client to provide all final text, images, branding assets (logos, colors) in timely manner.",
        "Client to review and provide feedback on milestones within 3 days of submission.",
        "Post-launch support: 30 days of free bug fixing and minor adjustments included.",
        "Any additional features or major changes beyond agreed scope will be quoted separately.",
        "Upon full payment, all intellectual property rights transfer to the Client.",
        "Developer retains right to display completed project in professional portfolio.",
        "Both parties agree to maintain confidentiality of sensitive business information.",
    ]
    
    for i, term in enumerate(terms, 1):
        doc.add_paragraph(term, style='List Number')
    
    doc.add_paragraph()
    
    # Acceptance
    add_section_title(doc, "ACCEPTANCE & SIGNATURE")
    
    doc.add_paragraph("To accept this quotation, please sign below and return to the provider.")
    doc.add_paragraph()
    
    sig_table = doc.add_table(rows=3, cols=2)
    
    sig_table.rows[0].cells[0].text = "Accepted by (Client): _____________________"
    sig_table.rows[0].cells[1].text = "Date: _____________________"
    
    sig_table.rows[1].cells[0].text = "Name: _____________________"
    sig_table.rows[1].cells[1].text = "Title: _____________________"
    
    sig_table.rows[2].cells[0].text = "Signed by (Provider): _____________________"
    sig_table.rows[2].cells[1].text = "Date: _____________________"
    
    doc.add_paragraph()
    
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.add_run("Thank you for considering our services!")
    footer_para.runs[0].font.italic = True
    
    return doc

def create_development_agreement_v2():
    """Create Updated Website Development Agreement"""
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # Title
    add_header(doc, "WEBSITE DEVELOPMENT AGREEMENT", "YoNISeRD Project")
    
    # Date and Parties
    doc.add_paragraph("Effective Date: 01st February 2026")
    doc.add_paragraph()
    
    doc.add_heading("PARTIES", level=1).runs[0].font.color.rgb = RGBColor(5, 15, 42)
    
    parties_text = """This Website Development Agreement ("Agreement") is entered into between:

DEVELOPER:
Derrick Omwanza
Professional Web Developer & Software Engineer
Nairobi, Kenya
Email: derrick@example.com
Phone: +254 XXX XXX XXX
(Hereinafter referred to as the "Developer")

AND

CLIENT:
Youth Network Integrated Services for Research and Development (YoNISeRD)
NGO Organization focused on Youth Research and Development
Kisii, Kenya
Email: [Client Email]
Phone: [Client Phone]
(Hereinafter referred to as the "Client")"""
    
    doc.add_paragraph(parties_text)
    doc.add_paragraph()
    
    # Scope of Services
    add_section_title(doc, "1. SCOPE OF SERVICES")
    
    doc.add_heading("1.1 Website Design & Development", level=2)
    
    design_scope = """The Developer agrees to design and develop a professional, fully-responsive website for the Client using the following technologies and methodologies:

Technology Stack:
- Framework: React 19 with React Router 7 for navigation
- Styling: TailwindCSS for responsive design
- Build Tool: Create React App
- Deployment: Netlify hosting platform
- Version Control: Git/GitHub repository

Design Specifications:
- Responsive design (Mobile-first approach)
- Breakpoints: Mobile (<640px), Tablet (640-1024px), Desktop (>1024px)
- Modern, professional aesthetic
- Brand consistency throughout all pages
- Color Scheme: Primary blue (#7EBBBF), Dark (#050F2A), Purple (#B8A0FF)"""
    
    doc.add_paragraph(design_scope)
    doc.add_paragraph()
    
    doc.add_heading("1.2 Website Pages & Features", level=2)
    
    pages_details = {
        "Home Page": "Hero banner with CTA, impact statistics, program highlights, latest stories, core values, CTA section",
        "About Page": "Mission/Vision statements, organization details, CEO message with portrait, team profiles (5 members), organizational timeline (6 milestones), awards/recognition, testimonials, values",
        "Programs Page": "Overview of 5 flagship programs, detailed descriptions, target audiences, expected outcomes, partnership opportunities",
        "Our Work Page": "Gallery of 9 community initiatives, story cards with tags, filtering capability, story modal for detailed view",
        "Gallery Page": "Standalone image gallery (6+ images), auto-rotating slideshow, manual navigation, responsive design, touch/swipe support",
        "News & Stories Page": "8+ stories with category filtering, story modal for detailed view, publication dates, author information, search functionality",
        "Impact Page": "Animated impact metrics, statistics visualization, beneficiary testimonials, photo timeline, data-driven narrative",
        "Contact Page": "Contact form with validation, form submission handling, contact information, volunteer CTA, partnership inquiry section",
    }
    
    for page, description in pages_details.items():
        para = doc.add_paragraph()
        para.add_run(f"{page}: ").font.bold = True
        para.add_run(description)
    
    doc.add_paragraph()
    
    doc.add_heading("1.3 Technical Features & Components", level=2)
    
    technical_features = [
        "Responsive Navigation System: Desktop dropdowns (3 menus × 3-4 items each), mobile hamburger menu, sticky header",
        "Contact Form: Client-side validation, error messages, success feedback, Netlify Forms integration, email notifications",
        "Image Slideshow Component: Auto-rotation, manual controls (prev/next), touch/swipe support, responsive aspect ratio (16:9)",
        "Impact Metrics Component: Scroll-triggered animations, animated counters, responsive layout",
        "Story Modal Component: Detailed story view, smooth transitions, close functionality",
        "ErrorBoundary Component: Graceful error handling and user-friendly error messages",
        "Footer: Social media links (LinkedIn, Twitter, Facebook, Instagram), quick navigation links, contact information",
        "Animations & Transitions: Smooth hover effects, fade-in animations, slide-up effects, respects reduced motion preferences",
        "Search Engine Optimization: Meta tags, Open Graph, Twitter Cards, Schema.org structured data, sitemap",
        "Accessibility Features: WCAG 2.1 Level AA compliance, skip links, ARIA labels, keyboard navigation, proper heading hierarchy",
        "Performance Optimization: CSS variable consolidation, GPU acceleration, will-change properties, CSS containment, lazy loading",
    ]
    
    for feature in technical_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_paragraph()
    
    # Project Timeline
    add_section_title(doc, "2. PROJECT TIMELINE & MILESTONES")
    
    timeline_details = """Project Commencement Date: [To be mutually agreed]
Estimated Completion Date: [Project already complete - ready for deployment]

Project Phases:
1. Design & Prototype - [COMPLETED]
2. Development & Implementation - [COMPLETED]
3. Quality Assurance & Testing - [COMPLETED with 100/100 score]
4. Final Optimization & Refinement - [COMPLETED]
5. Launch & Deployment - [READY FOR DEPLOYMENT]

The website has been fully developed, comprehensively tested, and optimized for production deployment. All milestones have been successfully completed."""
    
    doc.add_paragraph(timeline_details)
    doc.add_paragraph()
    
    # Financial Terms
    add_section_title(doc, "3. FINANCIAL TERMS")
    
    doc.add_heading("3.1 Total Project Fee", level=2)
    
    financial_text = """The total fee for the complete website development project is:

KSh 16,500 (Sixteen Thousand Five Hundred Kenyan Shillings)

This fee includes:
- Complete website design and development
- All 8 pages with full functionality
- Component development (navigation, forms, slideshows, etc.)
- Quality assurance and testing
- Performance optimization
- Deployment configuration and setup
- 30 days of post-launch support and bug fixes"""
    
    doc.add_paragraph(financial_text)
    doc.add_paragraph()
    
    doc.add_heading("3.2 Payment Schedule", level=2)
    
    payment_schedule_table = doc.add_table(rows=3, cols=3)
    payment_schedule_table.style = 'Light Grid Accent 1'
    
    header = payment_schedule_table.rows[0].cells
    header[0].text = "Payment Milestone"
    header[1].text = "Amount (KSh)"
    header[2].text = "Due Date / Condition"
    for cell in header:
        cell.paragraphs[0].runs[0].font.bold = True
    
    payment_schedule_table.rows[1].cells[0].text = "Initial Deposit (50%)"
    payment_schedule_table.rows[1].cells[1].text = "8,250"
    payment_schedule_table.rows[1].cells[2].text = "Upon signing this Agreement"
    
    payment_schedule_table.rows[2].cells[0].text = "Final Payment (50%)"
    payment_schedule_table.rows[2].cells[1].text = "8,250"
    payment_schedule_table.rows[2].cells[2].text = "Upon successful deployment and handover"
    
    doc.add_paragraph()
    
    doc.add_heading("3.3 Payment Methods", level=2)
    
    payment_methods = """Payments shall be made via one of the following methods to:

M-PESA (Preferred):
- Name: Derrick Omwanza
- Phone: +254 XXX XXX XXX
- Till Number: [Till Number]

Bank Transfer:
- Bank Name: [Bank Name]
- Account Name: Derrick Omwanza
- Account Number: [Account Number]
- Branch: [Branch]
- Swift Code: [Swift Code]

Payment confirmation should be sent to the Developer via email with transaction reference."""
    
    doc.add_paragraph(payment_methods)
    doc.add_paragraph()
    
    # Client Responsibilities
    add_section_title(doc, "4. CLIENT RESPONSIBILITIES")
    
    responsibilities = [
        "Provide all necessary text, images, and branding assets (logos, color specifications) in timely manner",
        "Provide access to necessary accounts (domain registrar, hosting, email) if required",
        "Review and provide feedback on deliverables within 3 business days of submission",
        "Provide accurate organization information for website content",
        "Communicate any changes or new requirements promptly",
        "Test the website thoroughly after deployment",
        "Report any bugs or issues within the 30-day support period",
    ]
    
    for responsibility in responsibilities:
        doc.add_paragraph(responsibility, style='List Bullet')
    
    doc.add_paragraph()
    
    # Developer Responsibilities
    add_section_title(doc, "5. DEVELOPER RESPONSIBILITIES")
    
    dev_responsibilities = [
        "Deliver a fully-functional, professionally-designed website meeting all specifications",
        "Ensure the website is responsive and works on all modern browsers and devices",
        "Implement all requested features and functionality as specified",
        "Provide comprehensive testing and quality assurance (100/100 QA score achieved)",
        "Optimize website for performance, security, and accessibility (WCAG 2.1 AA compliance)",
        "Provide clear deployment and handover documentation",
        "Offer 30 days of free post-launch support for bug fixes and minor adjustments",
        "Provide source code and all necessary project files to the Client",
        "Maintain confidentiality of Client information and business data",
    ]
    
    for responsibility in dev_responsibilities:
        doc.add_paragraph(responsibility, style='List Bullet')
    
    doc.add_paragraph()
    
    # Intellectual Property
    add_section_title(doc, "6. INTELLECTUAL PROPERTY RIGHTS")
    
    ip_text = """6.1 Custom Code & Design:
Upon receipt of full payment, the Developer assigns to the Client all intellectual property rights, title, and interest in:
- Custom website design and layout
- Custom-developed code and components
- Graphics and visual elements created specifically for the project
- All project files, source code, and documentation

6.2 Developer Portfolio Rights:
The Developer retains the right to:
- Display the completed website in professional portfolio
- Reference YoNISeRD as a client in business materials
- Use the project as a case study for marketing purposes
- Discuss non-confidential project details in professional contexts

6.3 Third-Party Components:
Any third-party libraries, frameworks, or open-source code used remain subject to their original licenses. The Client agrees to comply with these licenses."""
    
    doc.add_paragraph(ip_text)
    doc.add_paragraph()
    
    # Confidentiality
    add_section_title(doc, "7. CONFIDENTIALITY & DATA PROTECTION")
    
    confidentiality_text = """Both parties agree to maintain strict confidentiality of all non-public information obtained during the course of this project, including:
- Donor lists and funding information
- Internal organizational strategies
- Financial data and budgets
- Staff information
- Proprietary methodologies and approaches
- Beneficiary personal information

The Developer agrees to implement appropriate security measures to protect Client data and will not share any Client information with third parties without written consent."""
    
    doc.add_paragraph(confidentiality_text)
    doc.add_paragraph()
    
    # Maintenance & Support
    add_section_title(doc, "8. POST-LAUNCH MAINTENANCE & SUPPORT")
    
    maintenance_text = """8.1 Free Support Period:
This Agreement includes 30 days of free post-launch support starting from the deployment date. During this period, the Developer will:
- Monitor website performance and functionality
- Fix any bugs or technical issues at no additional cost
- Provide minor updates and adjustments
- Assist with content updates and basic maintenance
- Provide technical support via email

8.2 Beyond Support Period:
After the 30-day support period, additional maintenance and development will be billed separately at the rate of KSh 1,500 per day for:
- Major feature additions or modifications
- Performance optimization and tuning
- Security updates and patches
- Content management and updates
- Technical support and troubleshooting

8.3 Hosting & Domain:
The Client is responsible for:
- Maintaining hosting account with Netlify (or alternative provider)
- Managing domain registration and renewal
- Keeping SSL certificate current and valid
- Managing email accounts and configurations"""
    
    doc.add_paragraph(maintenance_text)
    doc.add_paragraph()
    
    # Limitation of Liability
    add_section_title(doc, "9. LIMITATION OF LIABILITY")
    
    liability_text = """The Developer's total liability under this Agreement shall not exceed the total amount paid by the Client. The Developer shall not be liable for:
- Any indirect, incidental, or consequential damages
- Loss of data or business interruption
- Third-party claims or actions
- Client's failure to provide timely information or access

The Client assumes responsibility for:
- Backing up website data and content
- Regular security monitoring and updates (after support period)
- Compliance with applicable laws and regulations"""
    
    doc.add_paragraph(liability_text)
    doc.add_paragraph()
    
    # Termination
    add_section_title(doc, "10. TERMINATION & DISPUTE RESOLUTION")
    
    termination_text = """10.1 Termination for Cause:
Either party may terminate this Agreement immediately upon written notice if the other party materially breaches a term and fails to cure the breach within 7 days of written notice.

10.2 Termination for Convenience:
The Client may terminate this project prior to completion, but must pay the Developer for:
- All work completed to date
- Non-refundable deposits
- Reasonable compensation for preparations and planning

10.3 Effect of Termination:
Upon termination:
- The Developer will deliver all completed work and files
- The Client must pay all outstanding invoices
- Both parties release all future claims
- Confidentiality obligations remain in effect

10.4 Dispute Resolution:
Any disputes arising from this Agreement shall be resolved through:
1. Good-faith negotiation between parties
2. Mediation (if negotiation fails)
3. Arbitration or legal proceedings (if necessary)"""
    
    doc.add_paragraph(termination_text)
    doc.add_paragraph()
    
    # General Terms
    add_section_title(doc, "11. GENERAL TERMS & CONDITIONS")
    
    general_terms = [
        "This Agreement constitutes the entire agreement between the parties and supersedes all prior discussions.",
        "Any modifications must be made in writing and signed by both parties.",
        "This Agreement shall be governed by the laws of Kenya.",
        "The Client and Developer are independent contractors; no employment relationship is created.",
        "If any provision is found invalid, remaining provisions continue in effect.",
        "The Developer is not responsible for third-party platform failures or changes.",
        "The Client agrees to provide truthful and accurate information for the website.",
        "Website content remains the Client's responsibility regarding accuracy and legality.",
    ]
    
    for i, term in enumerate(general_terms, 1):
        doc.add_paragraph(term, style='List Number')
    
    doc.add_paragraph()
    
    # Signatures
    add_section_title(doc, "12. SIGNATURES & ACCEPTANCE")
    
    doc.add_paragraph("By signing below, both parties acknowledge they have read, understood, and agree to all terms and conditions of this Agreement.")
    doc.add_paragraph()
    
    sig_section = doc.add_paragraph()
    sig_section.add_run("DEVELOPER:\n").font.bold = True
    sig_section.add_run("\nSignature: _____________________\n")
    sig_section.add_run("Name: Derrick Omwanza\n")
    sig_section.add_run("Date: _____________________\n")
    
    doc.add_paragraph()
    
    sig_section2 = doc.add_paragraph()
    sig_section2.add_run("CLIENT (YoNISeRD):\n").font.bold = True
    sig_section2.add_run("\nSignature: _____________________\n")
    sig_section2.add_run("Name: _____________________\n")
    sig_section2.add_run("Title: _____________________\n")
    sig_section2.add_run("Date: _____________________\n")
    
    doc.add_paragraph()
    
    # Footer
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("Website Development Agreement - YoNISeRD Project")
    footer_run.font.size = Pt(9)
    footer_run.font.italic = True
    footer_run.font.color.rgb = RGBColor(126, 187, 191)
    
    return doc

def main():
    """Generate all professional documents"""
    
    print("[*] Generating Professional Business Documents...\n")
    
    try:
        # Generate Invoice 2
        print("[*] Creating INVOICE_2_YoNISeRD.docx...")
        invoice_doc = create_invoice_v2()
        invoice_doc.save('c:/Users/ADMIN/ynis-rd-website/INVOICE_2_YoNISeRD.docx')
        print("[OK] INVOICE_2_YoNISeRD.docx created successfully\n")
        
        # Generate Project Quotation
        print("[*] Creating PROJECT_QUOTATION_V2_YoNISeRD.docx...")
        quotation_doc = create_project_quotation_v2()
        quotation_doc.save('c:/Users/ADMIN/ynis-rd-website/PROJECT_QUOTATION_V2_YoNISeRD.docx')
        print("[OK] PROJECT_QUOTATION_V2_YoNISeRD.docx created successfully\n")
        
        # Generate Development Agreement
        print("[*] Creating WEBSITE_DEVELOPMENT_AGREEMENT_V2_YoNISeRD.docx...")
        agreement_doc = create_development_agreement_v2()
        agreement_doc.save('c:/Users/ADMIN/ynis-rd-website/WEBSITE_DEVELOPMENT_AGREEMENT_V2_YoNISeRD.docx')
        print("[OK] WEBSITE_DEVELOPMENT_AGREEMENT_V2_YoNISeRD.docx created successfully\n")
        
        print("[SUCCESS] All documents generated successfully!\n")
        print("Documents created:")
        print("  1. INVOICE_2_YoNISeRD.docx")
        print("  2. PROJECT_QUOTATION_V2_YoNISeRD.docx")
        print("  3. WEBSITE_DEVELOPMENT_AGREEMENT_V2_YoNISeRD.docx")
        print("\nAll documents are ready for client delivery.")
        
    except Exception as e:
        print(f"[ERROR] Failed to generate documents: {e}")

if __name__ == '__main__':
    main()

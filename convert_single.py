from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def convert_md_to_docx(input_file, output_file):
    # Read the markdown file
    with open(input_file, 'r', encoding='utf-8') as f:
        content = f.read()

    # Create a new Word document
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Split content by lines
    lines = content.split('\n')

    for line in lines:
        stripped = line.strip()

        # Handle headings
        if stripped.startswith('# '):
            heading = stripped[2:].strip()
            p = doc.add_heading(heading, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif stripped.startswith('## '):
            heading = stripped[3:].strip()
            p = doc.add_heading(heading, level=2)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif stripped.startswith('### '):
            heading = stripped[4:].strip()
            p = doc.add_heading(heading, level=3)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif stripped.startswith('#### '):
            heading = stripped[5:].strip()
            p = doc.add_heading(heading, level=4)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Handle bullet points
        elif stripped.startswith('- '):
            text = stripped[2:].strip()
            p = doc.add_paragraph(text, style='List Bullet')

        # Handle horizontal rule
        elif stripped == '---':
            p = doc.add_paragraph('_' * 50)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Handle regular paragraphs
        elif stripped:
            p = doc.add_paragraph()
            # Handle bold text
            if '**' in stripped:
                parts = stripped.split('**')
                for i, part in enumerate(parts):
                    if i % 2 == 1:  # Bold parts
                        run = p.add_run(part)
                        run.bold = True
                    else:
                        p.add_run(part)
            else:
                p.add_run(stripped)

    # Save the document
    doc.save(output_file)
    print(f"Created: {output_file}")

# Convert the missing content analysis
convert_md_to_docx('MISSING_CONTENT_ANALYSIS.md', 'MISSING_CONTENT_ANALYSIS.docx')

#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Convert Markdown documents to Word (.docx) format for client delivery
"""

import os
import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# Set output encoding to UTF-8
if sys.stdout.encoding != 'utf-8':
    sys.stdout = open(sys.stdout.fileno(), mode='w', encoding='utf-8', buffering=1)

def md_to_docx(md_file, docx_file):
    """Convert markdown file to Word document"""
    
    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create Document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # Process markdown content line by line
    lines = content.split('\n')
    
    for i, line in enumerate(lines):
        stripped = line.strip()
        
        # Skip empty lines
        if not stripped:
            doc.add_paragraph()
            continue
        
        # Handle headings
        if stripped.startswith('# '):
            # H1
            p = doc.add_heading(stripped[2:].strip(), level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        elif stripped.startswith('## '):
            # H2
            doc.add_heading(stripped[3:].strip(), level=2)
            continue
        elif stripped.startswith('### '):
            # H3
            doc.add_heading(stripped[4:].strip(), level=3)
            continue
        elif stripped.startswith('#### '):
            # H4
            doc.add_heading(stripped[5:].strip(), level=4)
            continue
        
        # Handle tables (markdown format: | col | col |)
        if '|' in stripped:
            # Skip table separator lines
            if all(c in '-| ' for c in stripped):
                continue
            
            # Parse table row
            cells = [cell.strip() for cell in stripped.split('|')[1:-1]]
            
            # If this is the first row with content, create table
            if i == 0 or (i > 0 and '|' not in lines[i-1].strip()):
                num_cols = len(cells)
                table = doc.add_table(rows=1, cols=num_cols)
                table.style = 'Light Grid Accent 1'
                
                # Add header row
                header_cells = table.rows[0].cells
                for j, cell_text in enumerate(cells):
                    header_cells[j].text = cell_text
                    # Bold header
                    for paragraph in header_cells[j].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
            else:
                # Add data row to existing table
                if 'table' in locals():
                    row_cells = table.add_row().cells
                    for j, cell_text in enumerate(cells):
                        row_cells[j].text = cell_text
            continue
        
        # Handle code blocks
        if stripped.startswith('```'):
            # Skip code block markers
            continue
        
        # Handle bold text **text**
        if '**' in stripped:
            p = doc.add_paragraph()
            parts = stripped.split('**')
            for j, part in enumerate(parts):
                if j % 2 == 0:
                    # Regular text
                    if part:
                        p.add_run(part)
                else:
                    # Bold text
                    run = p.add_run(part)
                    run.font.bold = True
            continue
        
        # Handle checkmarks and special characters
        if stripped.startswith('✅'):
            p = doc.add_paragraph(stripped, style='List Bullet')
        elif stripped.startswith('❌'):
            p = doc.add_paragraph(stripped, style='List Bullet')
        elif stripped.startswith('🟢'):
            p = doc.add_paragraph(stripped)
        elif stripped.startswith('- '):
            # Bullet point
            doc.add_paragraph(stripped[2:], style='List Bullet')
        elif stripped.startswith('* '):
            # Bullet point
            doc.add_paragraph(stripped[2:], style='List Bullet')
        elif re.match(r'^\d+\. ', stripped):
            # Numbered list
            doc.add_paragraph(stripped[3:], style='List Number')
        else:
            # Regular paragraph
            doc.add_paragraph(stripped)
    
    # Save document
    doc.save(docx_file)

def main():
    """Convert all important client documents"""
    
    documents = [
        'QA_EXECUTIVE_SUMMARY.md',
        'COMPREHENSIVE_QA_AUDIT_REPORT.md',
        'FINAL_PRODUCTION_STATUS.md',
        'CSS_OPTIMIZATION_SUMMARY.md',
    ]
    
    base_path = Path('c:/Users/ADMIN/ynis-rd-website')
    
    print("[*] Converting Markdown documents to Word format...\n")
    
    for md_file in documents:
        md_path = base_path / md_file
        docx_path = base_path / md_file.replace('.md', '.docx')
        
        if md_path.exists():
            try:
                md_to_docx(str(md_path), str(docx_path))
                print(f"[OK] Converted: {md_file}")
            except Exception as e:
                print(f"[ERROR] Error converting {md_file}: {e}")
        else:
            print(f"[WARN] File not found: {md_file}")
    
    print("\n[SUCCESS] Conversion complete!")
    print("\nGenerated Word documents:")
    print("  - QA_EXECUTIVE_SUMMARY.docx")
    print("  - COMPREHENSIVE_QA_AUDIT_REPORT.docx")
    print("  - FINAL_PRODUCTION_STATUS.docx")
    print("  - CSS_OPTIMIZATION_SUMMARY.docx")

if __name__ == '__main__':
    main()

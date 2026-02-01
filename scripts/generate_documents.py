from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os

def convert_md_to_docx(input_file, output_file):
    if not os.path.exists(input_file):
        print(f"File not found: {input_file}")
        return

    # Read the markdown file
    with open(input_file, 'r', encoding='utf-8') as f:
        content = f.read()

    # Create a new Word document
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Split content by lines
    lines = content.split('\n')
    
    # Track if we are inside a table
    in_table = False
    table_data = []

    for line in lines:
        stripped = line.strip()
        
        # Handle Tables (Basic support for markdown tables)
        if stripped.startswith('|'):
            in_table = True
            # Parse table row
            cells = [c.strip() for c in stripped.split('|') if c.strip()]
            # Skip separator lines like |---|---|
            if '---' in stripped:
                continue
            table_data.append(cells)
            continue
        elif in_table:
            # End of table, render it
            if table_data:
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Table Grid'
                for i, row in enumerate(table_data):
                    for j, cell_text in enumerate(row):
                        # Handle <br> in table cells
                        cell_text = cell_text.replace('<br>', '\n')
                        # Handle bold in table cells
                        if '**' in cell_text:
                            p = table.cell(i, j).paragraphs[0]
                            parts = cell_text.split('**')
                            for k, part in enumerate(parts):
                                if k % 2 == 1:
                                    run = p.add_run(part)
                                    run.bold = True
                                else:
                                    p.add_run(part)
                        else:
                            table.cell(i, j).text = cell_text
            in_table = False
            table_data = []
        
        # Handle headings
        if stripped.startswith('# '):
            heading = stripped[2:].strip()
            p = doc.add_heading(heading, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif stripped.startswith('## '):
            heading = stripped[3:].strip()
            p = doc.add_heading(heading, level=2)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif stripped.startswith('### '):
            heading = stripped[4:].strip()
            p = doc.add_heading(heading, level=3)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif stripped.startswith('#### '):
            heading = stripped[5:].strip()
            p = doc.add_heading(heading, level=4)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Handle checkboxes and bullet points
        elif stripped.startswith('- [ ]'):
            text = stripped[5:].strip()
            p = doc.add_paragraph('☐ ' + text, style='List Bullet')
        elif stripped.startswith('- '):
            text = stripped[2:].strip()
            p = doc.add_paragraph(text, style='List Bullet')
        elif stripped.startswith('* '):
            text = stripped[2:].strip()
            p = doc.add_paragraph(text, style='List Bullet')
        
        # Handle numbered lists
        elif re.match(r'^[0-9]+\.\s', stripped):
            text = re.sub(r'^[0-9]+\.\s', '', stripped)
            p = doc.add_paragraph(text, style='List Number')
        
        # Handle horizontal rule
        elif stripped == '---':
            p = doc.add_paragraph('_' * 50)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Handle regular paragraphs and empty lines
        elif stripped:
            p = doc.add_paragraph()
            # Handle bold text
            if '**' in stripped:
                parts = stripped.split('**')
                for i, part in enumerate(parts):
                    if i % 2 == 1:  # Bold parts
                        run = p.add_run(part)
                        run.bold = True
                    else:
                        p.add_run(part)
            else:
                p.add_run(stripped)
        else:
            p = doc.add_paragraph()

    # Save the document
    doc.save(output_file)
    print(f"Created: {output_file}")

# List of files to convert
files_to_convert = [
    ('WEBSITE_DEVELOPMENT_AGREEMENT.md', 'WEBSITE_DEVELOPMENT_AGREEMENT.docx'),
    ('PROJECT_QUOTATION.md', 'PROJECT_QUOTATION.docx'),
    ('INVOICE.md', 'INVOICE.docx'),
    ('CONTENT_CHECKLIST.md', 'CONTENT_CHECKLIST.docx'),
    ('MISSING_CONTENT_ANALYSIS.md', 'MISSING_CONTENT_ANALYSIS.docx')
]

for input_file, output_file in files_to_convert:
    convert_md_to_docx(input_file, output_file)

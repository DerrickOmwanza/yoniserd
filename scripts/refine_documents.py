#!/usr/bin/env python3
"""
Script to refine the three Word documents with feedback suggestions
"""

from docx import Document
from copy import deepcopy

def add_paragraph_after(doc, insert_index, text, style=None):
    """Helper to insert a new paragraph after specified index"""
    if insert_index < len(doc.paragraphs):
        elem = doc.paragraphs[insert_index]._element
        new_elem = deepcopy(elem)
        elem.addnext(new_elem)
        new_para = doc.paragraphs[insert_index + 1]
        new_para.text = text
        if style:
            new_para.style = style
        return new_para
    return None

def refine_invoice():
    """Refine INVOICE_2_YoNISeRD.docx"""
    doc = Document('INVOICE_2_YoNISeRD.docx')
    
    # Find and enhance total amount section
    for i, para in enumerate(doc.paragraphs):
        if 'Total Amount' in para.text:
            add_paragraph_after(doc, i, "Deposit Received: KSh 2,000")
            add_paragraph_after(doc, i+1, "Balance Due: KSh 14,500")
            break
    
    # Update status
    for para in doc.paragraphs:
        if 'Status:' in para.text:
            para.text = "Status: Ready for Payment"
    
    doc.save('INVOICE_2_YoNISeRD.docx')
    print("[OK] Invoice refined")

def refine_quotation():
    """Refine PROJECT_QUOTATION_V2_YoNISeRD.docx"""
    doc = Document('PROJECT_QUOTATION_V2_YoNISeRD.docx')
    
    # Update timeline dates
    date_updates = {
        'Launch & Deployment': 'Launch & Deployment - 2nd February 2026 (Completed)',
        'Post-Launch Support': 'Post-Launch Support - 30 days from 2nd February 2026'
    }
    
    for para in doc.paragraphs:
        for key, value in date_updates.items():
            if key in para.text:
                para.text = value
    
    # Find scope section and add change requests
    for i, para in enumerate(doc.paragraphs):
        if 'Scope of Work' in para.text:
            # Find Timeline section
            for j in range(i+1, len(doc.paragraphs)):
                if 'Timeline' in doc.paragraphs[j].text:
                    elem = doc.paragraphs[j]._element
                    heading_elem = deepcopy(elem)
                    clause_elem = deepcopy(elem)
                    elem.addprevious(heading_elem)
                    elem.addprevious(clause_elem)
                    
                    doc.paragraphs[j].text = "Change Requests"
                    doc.paragraphs[j].style = 'Heading 3'
                    doc.paragraphs[j+1].text = "Any changes beyond this scope will be billed separately."
                    break
            break
    
    doc.save('PROJECT_QUOTATION_V2_YoNISeRD.docx')
    print("[OK] Quotation refined")

def refine_agreement():
    """Refine WEBSITE_DEVELOPMENT_AGREEMENT_V2_YoNISeRD.docx"""
    doc = Document('WEBSITE_DEVELOPMENT_AGREEMENT_V2_YoNISeRD.docx')
    
    # Add source code handover clause
    for i, para in enumerate(doc.paragraphs):
        if 'Payment Terms' in para.text:
            for j in range(i+1, len(doc.paragraphs)):
                if 'Post-Launch Support' in doc.paragraphs[j].text:
                    add_paragraph_after(doc, j, "Source Code Handover", 'Heading 3')
                    add_paragraph_after(doc, j+1, 
                        "Complete source code will be provided only after full payment.")
                    break
            break
    
    # Add non-refundable deposit clause
    for i, para in enumerate(doc.paragraphs):
        if 'Termination' in para.text:
            for j in range(i+1, min(i+8, len(doc.paragraphs))):
                if 'Client' in doc.paragraphs[j].text and 'terminate' in doc.paragraphs[j].text.lower():
                    add_paragraph_after(doc, j, 
                        "Deposit Policy: The initial deposit of KSh 2,000 is non-refundable.")
                    break
            break
    
    doc.save('WEBSITE_DEVELOPMENT_AGREEMENT_V2_YoNISeRD.docx')
    print("[OK] Agreement refined")

if __name__ == '__main__':
    print("Refining all three Word documents...\n")
    try:
        refine_invoice()
        refine_quotation()
        refine_agreement()
        print("\n[SUCCESS] All documents refined successfully!")
    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()

#!/usr/bin/env python3
"""
Script to refine the three Word documents with feedback suggestions
"""

from docx import Document
from docx.shared import Pt

def refine_invoice():
    """Refine INVOICE_2_YoNISeRD.docx"""
    doc = Document('INVOICE_2_YoNISeRD.docx')
    
    # Find the services/amount section and add payment breakdown
    modified = False
    for i, para in enumerate(doc.paragraphs):
        if 'Total Amount' in para.text and '14,000' in para.text and not modified:
            # Add blank line
            new_para = doc.paragraphs[i]._element
            parent = new_para.getparent()
            
            # Create new paragraphs for deposit and balance
            from copy import deepcopy
            
            # Blank line
            blank = deepcopy(new_para)
            parent.insert(parent.index(new_para) + 1, blank)
            
            # Deposit line
            deposit_elem = deepcopy(new_para)
            deposit_elem.clear_content()
            parent.insert(parent.index(blank) + 1, deposit_elem)
            
            # Balance line
            balance_elem = deepcopy(new_para)
            balance_elem.clear_content()
            parent.insert(parent.index(deposit_elem) + 1, balance_elem)
            
            # Add text to new paragraphs
            doc.paragraphs[i+2].text = "Deposit Received: KSh 2,000"
            doc.paragraphs[i+3].text = "Balance Due: KSh 14,500"
            
            modified = True
    
    # Find and enhance status section
    for para in doc.paragraphs:
        if 'Status:' in para.text or 'Ready for Payment' in para.text:
            para.text = "Status: Ready for Payment"
    
    # Add note about partial payments
    for i, para in enumerate(doc.paragraphs):
        if 'Payment Instructions' in para.text:
            # Insert note before payment instructions
            new_elem = para._element
            parent = new_elem.getparent()
            
            from copy import deepcopy
            note_elem = deepcopy(new_elem)
            parent.insert(parent.index(new_elem), note_elem)
            
            note_para = doc.paragraphs[i]
            note_para.text = "Note: Partial payments do not constitute full settlement. Full payment is required."
            break
    
    doc.save('INVOICE_2_YoNISeRD.docx')
    print("[OK] Invoice refined")

def refine_quotation():
    """Refine PROJECT_QUOTATION_V2_YoNISeRD.docx"""
    doc = Document('PROJECT_QUOTATION_V2_YoNISeRD.docx')
    
    # Update timeline with actual dates
    for para in doc.paragraphs:
        if 'Launch & Deployment' in para.text:
            para.text = "Launch & Deployment - 2nd February 2026 (Completed)"
        elif 'Post-Launch Support' in para.text and 'days' in para.text.lower():
            para.text = "Post-Launch Support - 30 days from 2nd February 2026"
    
    # Find scope section and add change requests clause
    added_clause = False
    for i, para in enumerate(doc.paragraphs):
        if 'Scope of Work' in para.text and not added_clause:
            # Look for the end of scope items
            for j in range(i+1, len(doc.paragraphs)):
                # Find where to insert (before Timeline or similar)
                if 'Timeline' in doc.paragraphs[j].text or 'Payment' in doc.paragraphs[j].text:
                    # Insert Change Requests heading
                    insert_elem = doc.paragraphs[j]._element
                    parent = insert_elem.getparent()
                    
                    from copy import deepcopy
                    heading_elem = deepcopy(insert_elem)
                    parent.insert(parent.index(insert_elem), heading_elem)
                    
                    clause_elem = deepcopy(insert_elem)
                    parent.insert(parent.index(heading_elem) + 1, clause_elem)
                    
                    new_heading = doc.paragraphs[j]
                    new_heading.text = "Change Requests"
                    new_heading.style = 'Heading 3'
                    
                    new_clause = doc.paragraphs[j+1]
                    new_clause.text = "Any changes beyond this defined scope will be billed separately at agreed rates."
                    
                    added_clause = True
                    break
    
    doc.save('PROJECT_QUOTATION_V2_YoNISeRD.docx')
    print("[OK] Quotation refined")

def refine_agreement():
    """Refine WEBSITE_DEVELOPMENT_AGREEMENT_V2_YoNISeRD.docx"""
    doc = Document('WEBSITE_DEVELOPMENT_AGREEMENT_V2_YoNISeRD.docx')
    
    # Find Payment Terms and enhance them
    for i, para in enumerate(doc.paragraphs):
        if 'Payment Terms' in para.text:
            # Look ahead to find where to add source code handover
            for j in range(i+1, min(i+10, len(doc.paragraphs))):
                if 'Post-Launch Support' in doc.paragraphs[j].text or 'Support Period' in doc.paragraphs[j].text:
                    # Insert before post-launch support
                    insert_elem = doc.paragraphs[j]._element
                    parent = insert_elem.getparent()
                    
                    from copy import deepcopy
                    heading_elem = deepcopy(insert_elem)
                    parent.insert(parent.index(insert_elem), heading_elem)
                    
                    content_elem = deepcopy(insert_elem)
                    parent.insert(parent.index(heading_elem) + 1, content_elem)
                    
                    new_heading = doc.paragraphs[j]
                    new_heading.text = "Source Code Handover"
                    new_heading.style = 'Heading 3'
                    
                    new_content = doc.paragraphs[j+1]
                    new_content.text = "Complete source code and project files will be provided only after full payment (KSh 16,500) is received."
                    
                    break
    
    # Find Termination clause and add non-refundable deposit note
    for i, para in enumerate(doc.paragraphs):
        if 'Termination' in para.text:
            for j in range(i+1, min(i+8, len(doc.paragraphs))):
                if 'Client' in doc.paragraphs[j].text and 'terminate' in doc.paragraphs[j].text.lower():
                    # Add non-refundable note after this paragraph
                    insert_elem = doc.paragraphs[j]._element
                    parent = insert_elem.getparent()
                    
                    from copy import deepcopy
                    note_elem = deepcopy(insert_elem)
                    parent.insert(parent.index(insert_elem) + 1, note_elem)
                    
                    new_note = doc.paragraphs[j+1]
                    new_note.text = "Deposit Policy: The initial deposit of KSh 2,000 is non-refundable."
                    break
    
    doc.save('WEBSITE_DEVELOPMENT_AGREEMENT_V2_YoNISeRD.docx')
    print("[OK] Agreement refined")

if __name__ == '__main__':
    print("Refining all three Word documents...\n")
    
    try:
        refine_invoice()
        refine_quotation()
        refine_agreement()
        print("\n[SUCCESS] All documents refined successfully!")
    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()
